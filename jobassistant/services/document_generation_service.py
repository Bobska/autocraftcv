"""
Document generation service for creating PDF and DOCX files
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import os
from typing import Dict, Tuple
from django.core.files.base import ContentFile
from django.conf import settings
import logging

logger = logging.getLogger(__name__)


class DocumentGenerationService:
    """Service for generating PDF and DOCX documents"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation"""
        
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=20,
        )
        
        # Header style
        self.header_style = ParagraphStyle(
            'CustomHeader',
            parent=self.styles['Heading2'],
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.darkblue,
            spaceBefore=15,
            spaceAfter=10,
        )
        
        # Body style
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            leftIndent=0,
        )
        
        # Contact info style
        self.contact_style = ParagraphStyle(
            'ContactStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            alignment=TA_CENTER,
            spaceAfter=15,
        )
    
    def generate_pdf(self, content: str, document_type: str, user_name: str = "") -> ContentFile:
        """
        Generate PDF document from content
        Returns Django ContentFile
        """
        buffer = io.BytesIO()
        
        try:
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Build the document
            story = []
            
            if document_type == 'cover_letter':
                story = self._build_cover_letter_pdf(content, user_name)
            elif document_type == 'resume':
                story = self._build_resume_pdf(content, user_name)
            else:
                # Generic document
                story = self._build_generic_pdf(content)
            
            doc.build(story)
            
            buffer.seek(0)
            filename = f"{document_type}_{user_name.replace(' ', '_')}.pdf" if user_name else f"{document_type}.pdf"
            
            return ContentFile(buffer.getvalue(), name=filename)
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise e
        finally:
            buffer.close()
    
    def generate_docx(self, content: str, document_type: str, user_name: str = "") -> ContentFile:
        """
        Generate DOCX document from content
        Returns Django ContentFile
        """
        try:
            doc = Document()
            
            if document_type == 'cover_letter':
                self._build_cover_letter_docx(doc, content, user_name)
            elif document_type == 'resume':
                self._build_resume_docx(doc, content, user_name)
            else:
                # Generic document
                self._build_generic_docx(doc, content)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            filename = f"{document_type}_{user_name.replace(' ', '_')}.docx" if user_name else f"{document_type}.docx"
            
            return ContentFile(buffer.getvalue(), name=filename)
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise e
    
    def _build_cover_letter_pdf(self, content: str, user_name: str) -> list:
        """Build PDF story for cover letter"""
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue
            
            # First paragraph might be the header/name
            if i == 0 and user_name and user_name.lower() in para.lower():
                story.append(Paragraph(para, self.title_style))
                story.append(Spacer(1, 20))
            else:
                # Regular paragraph
                story.append(Paragraph(para, self.body_style))
                story.append(Spacer(1, 12))
        
        return story
    
    def _build_resume_pdf(self, content: str, user_name: str) -> list:
        """Build PDF story for resume"""
        story = []
        
        lines = content.split('\n')
        current_section = []
        section_title = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Detect section headers (all caps or specific patterns)
            if self._is_section_header(line):
                # Add previous section
                if current_section:
                    story.extend(self._format_resume_section(section_title, current_section))
                    current_section = []
                
                section_title = line
                story.append(Paragraph(line, self.header_style))
                story.append(Spacer(1, 8))
            
            # Name/contact info (first few lines)
            elif len(story) == 0 and (user_name and user_name.lower() in line.lower() or '@' in line or '(' in line):
                story.append(Paragraph(line, self.contact_style))
            
            else:
                current_section.append(line)
        
        # Add final section
        if current_section:
            story.extend(self._format_resume_section(section_title, current_section))
        
        return story
    
    def _build_generic_pdf(self, content: str) -> list:
        """Build generic PDF story"""
        story = []
        
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if para:
                story.append(Paragraph(para, self.body_style))
                story.append(Spacer(1, 12))
        
        return story
    
    def _build_cover_letter_docx(self, doc: Document, content: str, user_name: str):
        """Build DOCX document for cover letter"""
        
        # Set document margins
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        paragraphs = content.split('\n\n')
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue
            
            p = doc.add_paragraph(para)
            
            # First paragraph styling (if it contains name)
            if i == 0 and user_name and user_name.lower() in para.lower():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.bold = True
                run.font.size = Inches(0.2)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _build_resume_docx(self, doc: Document, content: str, user_name: str):
        """Build DOCX document for resume"""
        
        # Set document margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Section headers
            if self._is_section_header(line):
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.bold = True
                run.font.size = Inches(0.15)
                p.space_after = Inches(0.1)
            
            # Name/contact info
            elif user_name and user_name.lower() in line.lower() or '@' in line:
                p = doc.add_paragraph(line)
                if user_name and user_name.lower() in line.lower():
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.runs[0]
                    run.bold = True
                    run.font.size = Inches(0.18)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Regular content
            else:
                doc.add_paragraph(line)
    
    def _build_generic_docx(self, doc: Document, content: str):
        """Build generic DOCX document"""
        
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    
    def _format_resume_section(self, section_title: str, content_lines: list) -> list:
        """Format a resume section for PDF"""
        story = []
        
        for line in content_lines:
            line = line.strip()
            if line:
                # Check if it's a sub-header (job title, degree, etc.)
                if self._is_sub_header(line):
                    sub_style = ParagraphStyle(
                        'SubHeader',
                        parent=self.body_style,
                        fontSize=11,
                        fontName='Helvetica-Bold',
                        spaceBefore=8,
                        spaceAfter=4,
                    )
                    story.append(Paragraph(line, sub_style))
                else:
                    story.append(Paragraph(line, self.body_style))
                
                story.append(Spacer(1, 6))
        
        return story
    
    def _is_section_header(self, line: str) -> bool:
        """Check if line is a section header"""
        line = line.strip()
        
        # Common resume section headers
        section_headers = [
            'PROFESSIONAL SUMMARY', 'PROFESSIONAL EXPERIENCE', 'TECHNICAL SKILLS',
            'EDUCATION', 'CERTIFICATIONS', 'ACHIEVEMENTS', 'CORE COMPETENCIES',
            'WORK EXPERIENCE', 'SKILLS', 'SUMMARY', 'EXPERIENCE', 'PROJECTS',
            'AWARDS', 'PUBLICATIONS', 'LANGUAGES', 'INTERESTS'
        ]
        
        # Check if it's all caps
        if line.isupper() and len(line) > 3:
            return True
        
        # Check if it's a known section header
        if line.upper() in section_headers:
            return True
        
        # Check if it ends with colon
        if line.endswith(':') and len(line.split()) <= 3:
            return True
        
        return False
    
    def _is_sub_header(self, line: str) -> bool:
        """Check if line is a sub-header (job title, company, etc.)"""
        
        # Look for patterns like "Job Title | Company | Dates"
        if '|' in line and len(line.split('|')) >= 2:
            return True
        
        # Look for patterns with dates
        if any(char.isdigit() for char in line) and ('20' in line or '19' in line):
            return True
        
        # Look for company/university indicators
        company_indicators = ['Inc.', 'Corp.', 'LLC', 'Ltd.', 'University', 'College', 'Institute']
        if any(indicator in line for indicator in company_indicators):
            return True
        
        return False
    
    def generate_both_formats(self, content: str, document_type: str, user_name: str = "") -> Tuple[ContentFile, ContentFile]:
        """
        Generate both PDF and DOCX formats
        Returns: (pdf_file, docx_file)
        """
        pdf_file = self.generate_pdf(content, document_type, user_name)
        docx_file = self.generate_docx(content, document_type, user_name)
        
        return pdf_file, docx_file
